import validators,streamlit as st
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.chains.summarize import load_summarize_chain
from langchain_community.document_loaders import YoutubeLoader,UnstructuredURLLoader
import docx
from io import BytesIO
### code

def generate_doc(content):
    doc = docx.Document()
    
    # Split the content into sections
    sections = content.split('**')
    
    for i, section in enumerate(sections):
        if i % 2 == 0:  # Even indexes are normal text or empty
            continue
        else:  # Odd indexes are headers or content
            # Add the header in bold
            p = doc.add_paragraph()
            p.add_run(section.strip()).bold = True
            
            # If there's content following this header, add it as bullet points
            if i + 1 < len(sections):
                content = sections[i + 1]
                bullet_points = content.split('*')
                for point in bullet_points:
                    if point.strip():
                        doc.add_paragraph(point.strip(), style='List Bullet')
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


## sstreamlit APP
st.set_page_config(page_title="LangChain: Summarize Text From YT or Website", page_icon="🦜")
st.title("🦜 LangChain: Summarize Text From YT or Website")
st.subheader('Summarize URL')



## Get the Groq API Key and url(YT or website)to be summarized
st.sidebar.title("Settings")
api_key=st.sidebar.text_input("Enter your Groq API Key:",type="password")

generic_url=st.text_input("URL",label_visibility="collapsed")

## Gemma Model USsing Groq API
llm =ChatGroq(groq_api_key=api_key, model_name="Gemma-7b-It")

prompt_template="""
Provide a summary of the following content in 300 words:
Content:{text}

"""
prompt=PromptTemplate(template=prompt_template,input_variables=["text"])

if st.button("Summarize the Content from YT or Website"):
    ## Validate all the inputs
    if not api_key.strip() or not generic_url.strip():
        st.error("Please provide the information to get started")
    elif not validators.url(generic_url):
        st.error("Please enter a valid Url. It can may be a YT video utl or website url")

    else:
        try:
            with st.spinner("Waiting..."):
                ## loading the website or yt video data
                if "youtube.com" in generic_url:
                    loader=YoutubeLoader.from_youtube_url(generic_url,add_video_info=True)
                else:
                    loader=UnstructuredURLLoader(urls=[generic_url],ssl_verify=False,
                                                 headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_5_1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/116.0.0.0 Safari/537.36"})
                docs=loader.load()

                ## Chain For Summarization
                chain=load_summarize_chain(llm,chain_type="stuff",prompt=prompt)
                output_summary=chain.run(docs)
                st.success(output_summary)
                
                # Generate and offer download of .doc file
                doc_io = generate_doc(output_summary)
                st.download_button(
                    label="Download Response As docx file",
                    data=doc_io,
                    file_name="app_response.doc",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                
        except Exception as e:
            st.exception(f"Exception:{e}")
                    
